from docx import Document
import glob
import os
import json

os.chdir(r'D:\develop\code\CareerAgent\md')

# 查找文件
template_file = glob.glob('2023*.docx')[0]
v3_file = '职引未来_技术文档_v3.0.docx'

def extract_doc_structure(doc, name):
    """提取文档结构"""
    result = {
        'name': name,
        'paragraphs': [],
        'tables': []
    }

    # 提取段落
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else 'none'
            result['paragraphs'].append({
                'index': i,
                'style': style,
                'text': para.text
            })

    # 提取表格
    for idx, table in enumerate(doc.tables):
        table_data = {'index': idx, 'rows': []}
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_data['rows'].append(cells)
        result['tables'].append(table_data)

    return result

# 读取两个文档
template_doc = Document(template_file)
v3_doc = Document(v3_file)

# 提取结构
template_struct = extract_doc_structure(template_doc, template_file)
v3_struct = extract_doc_structure(v3_doc, v3_file)

# 保存为 JSON
with open('template_structure.json', 'w', encoding='utf-8') as f:
    json.dump(template_struct, f, ensure_ascii=False, indent=2)

with open('v3_structure.json', 'w', encoding='utf-8') as f:
    json.dump(v3_struct, f, ensure_ascii=False, indent=2)

# 打印模板的主要章节
print("=== 模板文件主要章节 ===")
styles_order = ['Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5']
for para in template_struct['paragraphs']:
    style = para['style']
    if any(s in style for s in ['Heading', '标题', 'Title', 'Subtitle', '目录']):
        print(f"{para['index']}: [{style}] {para['text']}")

# 打印模板的表格标题
print("\n=== 模板文件表格 ===")
for table in template_struct['tables']:
    if table['rows']:
        first_row = table['rows'][0] if table['rows'] else []
        print(f"表格 {table['index']+1}: {first_row[:5]}...")

print(f"\n模板文件共有 {len(template_struct['paragraphs'])} 段落，{len(template_struct['tables'])} 表格")
print(f"v3.0 文件共有 {len(v3_struct['paragraphs'])} 段落，{len(v3_struct['tables'])} 表格")
