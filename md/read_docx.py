from docx import Document
import os
import glob

os.chdir(r'D:\develop\code\CareerAgent\md')

# 查找文件
template_file = glob.glob('2023*.docx')[0]
print(f"模板文件：{template_file}")

# 读取模板
print("=== 读取模板文件 ===")
template_doc = Document(template_file)
print(f'模板段落数：{len(template_doc.paragraphs)}')
print(f'模板表格数：{len(template_doc.tables)}')

print('\n=== 模板标题结构 ===')
for i, para in enumerate(template_doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'none'
        if 'Heading' in style or '标题' in style or 'Title' in style or 'Subtitle' in style:
            print(f'{i}: [{style}] {para.text}')

# 读取 v3.0
print("\n\n=== 读取 v3.0 文件 ===")
v3_doc = Document('职引未来_技术文档_v3.0.docx')
print(f'v3.0 段落数：{len(v3_doc.paragraphs)}')
print(f'v3.0 表格数：{len(v3_doc.tables)}')

print('\n=== v3.0 标题结构 ===')
for i, para in enumerate(v3_doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'none'
        if 'Heading' in style or '标题' in style or 'Title' in style or 'Subtitle' in style:
            print(f'{i}: [{style}] {para.text}')

# 输出 v3.0 前言部分
print("\n\n=== v3.0 前言章节 ===")
in_intro = False
for i, para in enumerate(v3_doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'none'
        if 'Heading 1' in style and '前言' in para.text:
            in_intro = True
            print(f'\n{para.text}')
            continue
        if in_intro:
            if 'Heading 1' in style:
                break
            print(f'[{style}] {para.text}')

# 输出 v3.0 项目背景章节
print("\n\n=== v3.0 项目概述章节 ===")
in_overview = False
for i, para in enumerate(v3_doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'none'
        if 'Heading 1' in style and '项目概述' in para.text:
            in_overview = True
            print(f'\n{para.text}')
            continue
        if in_overview:
            if 'Heading 1' in style:
                break
            print(f'[{style}] {para.text[:200]}')
