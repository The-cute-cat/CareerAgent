#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
职引未来 - 技术文档 Word 生成脚本 (增强版)
按照参考文档格式进行美化，生成专业的技术文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False, color=None):
    """设置中文字体和样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, align='left'):
    """添加标题段落"""
    p = doc.add_paragraph()

    if level == 1:
        # 一级标题 - 黑体三号加粗
        set_chinese_font(p.add_run(text), '黑体', 16, bold=True)
        p.space_after = Pt(12)
        p.space_before = Pt(12)
    elif level == 2:
        # 二级标题 - 黑体四号加粗
        set_chinese_font(p.add_run(text), '黑体', 14, bold=True)
        p.space_after = Pt(10)
        p.space_before = Pt(10)
    elif level == 3:
        # 三级标题 - 黑体小四加粗
        set_chinese_font(p.add_run(text), '黑体', 12, bold=True)
        p.space_after = Pt(6)
        p.space_before = Pt(6)
    elif level == 4:
        # 四级标题 - 宋体小四加粗
        set_chinese_font(p.add_run(text), '宋体', 12, bold=True)
        p.space_after = Pt(6)

    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return p


def add_paragraph(doc, text, indent=True, line_spacing=1.5):
    """添加普通段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符
    p.paragraph_format.line_spacing = line_spacing
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 12)
    return p


def add_code_block(doc, code_text, language=''):
    """添加代码块"""
    # 添加灰色背景框模拟代码块
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.space_after = Pt(6)

    for line in code_text.split('\n'):
        run = p.add_run(line)
        set_chinese_font(run, 'Consolas', 9)
        run.font.color.rgb = RGBColor(40, 40, 40)

    return p


def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头样式
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 清空并设置表头
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_chinese_font(run, '黑体', 10, bold=True)
        p.paragraph_format.line_spacing = 1.0

    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            cell = row_cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # 清空并设置单元格内容
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(str(cell_text))
            set_chinese_font(run, '宋体', 9)

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    doc.add_paragraph()  # 表格后加空行
    return table


def create_document():
    """创建职引未来技术文档"""
    doc = Document()

    # 设置页面边距 (A4)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ==================== 封面 ====================
    add_heading(doc, "职引未来", level=1, align='center')
    add_heading(doc, "基于 AI 的大学生职业规划智能体系统", level=1, align='center')
    add_heading(doc, "技术设计文档", level=1, align='center')

    # 文档信息
    for _ in range(6):
        doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = info_para.add_run("版  本：v2.0\n")
    set_chinese_font(run, '宋体', 12)
    run = info_para.add_run("编制日期：2026 年 4 月 13 日\n")
    set_chinese_font(run, '宋体', 12)
    run = info_para.add_run("编制单位：职引未来开发团队\n")
    set_chinese_font(run, '宋体', 12)

    doc.add_page_break()

    # ==================== 修订记录 ====================
    add_heading(doc, "修订记录", level=1)

    revision_data = [
        ("版本", "修订日期", "修订内容", "编写人"),
        ("v1.0", "2026-03-01", "初稿完成", "开发团队"),
        ("v1.1", "2026-03-15", "补充数据库设计细节", "开发团队"),
        ("v1.2", "2026-03-25", "完善 API 接口文档", "开发团队"),
        ("v2.0", "2026-04-13", "全面更新技术架构与部署指南", "开发团队"),
    ]
    add_table(doc, ["版本", "修订日期", "修订内容", "编写人"], revision_data[1:], [2.5, 2.5, 6, 2])

    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading(doc, "目  录", level=1, align='center')

    toc_items = [
        ("第一章 项目概述", "1"),
        ("  1.1 系统简介", "1"),
        ("  1.2 核心功能模块", "2"),
        ("  1.3 技术栈概览", "3"),
        ("  1.4 项目结构", "4"),
        ("第二章 技术架构", "6"),
        ("  2.1 整体架构设计", "6"),
        ("  2.2 后端技术栈", "7"),
        ("  2.3 前端技术栈", "8"),
        ("  2.4 AI 服务技术栈", "9"),
        ("  2.5 核心模块设计", "10"),
        ("第三章 数据库设计", "12"),
        ("  3.1 数据库概述", "12"),
        ("  3.2 核心表设计", "13"),
        ("  3.3 索引优化建议", "20"),
        ("第四章 API 接口文档", "22"),
        ("  4.1 通用说明", "22"),
        ("  4.2 用户模块", "23"),
        ("  4.3 简历模块", "25"),
        ("  4.4 AI 智能助手模块", "26"),
        ("第五章 部署指南", "28"),
        ("  5.1 环境要求", "28"),
        ("  5.2 部署架构", "29"),
        ("  5.3 部署步骤", "30"),
        ("  5.4 监控与日志", "35"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(item + "\t\t\t第" + page + "页")
        set_chinese_font(run, '宋体', 11)

    doc.add_page_break()

    # ==================== 第一章 项目概述 ====================
    add_heading(doc, "第一章 项目概述", level=1)

    add_heading(doc, "1.1 系统简介", level=2)
    add_paragraph(doc,
        "「职引未来」是一个基于人工智能的大学生职业规划智能体系统，旨在帮助大学生进行科学的职业规划与就业准备。"
        "系统整合了职业测评、简历优化、面试准备、职业知识库、AI 智能问答等多维度功能，为用户提供全方位的职业发展支持。")
    add_paragraph(doc,
        "系统采用前后端分离的微服务架构设计，后端基于 Spring Boot 框架，AI 服务基于 FastAPI+Python 构建，"
        "通过 HTTP 接口进行通信，实现了业务逻辑与 AI 能力的解耦。")

    add_heading(doc, "1.2 核心功能模块", level=2)

    modules = [
        ("用户管理模块", "用户注册与登录、用户信息管理、会员体系管理、积分系统管理"),
        ("职业测评模块", "职业能力评估、技能雷达图展示、个人优势分析"),
        ("简历管理模块", "简历在线编辑、模板选择、解析与智能填充、导出功能"),
        ("AI 智能助手模块", "AI 职业规划问答、面试题智能生成、职业发展建议、用户画像分析"),
        ("面试管理模块", "面试题库管理、面试日历、面试回顾与反馈"),
        ("职位匹配模块", "职位搜索、智能岗位推荐、用户与岗位匹配分析"),
        ("知识库模块", "职业知识图谱、行业知识浏览、发展路径导航"),
    ]

    for i, (name, desc) in enumerate(modules, 1):
        add_paragraph(doc, f"（{i}）{name}：{desc}")

    add_heading(doc, "1.3 技术栈概览", level=2)

    tech_stack = [
        ("后端框架", "Spring Boot 2.7.x"),
        ("前端框架", "Vue 3 + TypeScript + Vite"),
        ("AI 服务", "FastAPI + LangChain + Python 3.12"),
        ("主数据库", "MySQL 8.0"),
        ("ORM 框架", "MyBatis-Plus 3.5.x"),
        ("缓存系统", "Redis 7.x"),
        ("消息队列", "RabbitMQ 3.12.x"),
        ("API 文档", "Swagger/OpenAPI 3.0"),
        ("AI 模型", "通义千问 (Qwen 3.5)"),
        ("向量数据库", "Milvus"),
        ("图数据库", "Neo4j/SQLite"),
    ]
    add_table(doc, ["组件类别", "技术选型"], tech_stack, [3.5, 7.5])

    add_heading(doc, "1.4 项目结构", level=2)

    project_structure = """CareerAgent/
├── career-planning-backend/          # Java 后端服务
│   ├── src/main/java/com/backend/
│   │   └── careerplanningbackend/
│   │       ├── config/              # 配置类 (安全、拦截器、跨域等)
│   │       ├── controller/          # 控制器层 (REST API)
│   │       ├── domain/              # 实体类与 DTO
│   │       ├── service/             # 业务逻辑层
│   │       ├── mapper/              # 数据访问层 (MyBatis)
│   │       ├── listener/            # 事件监听器
│   │       └── utils/               # 工具类
│   └── src/main/resources/
│       ├── mapper/                  # MyBatis XML 映射文件
│       ├── sql/                     # 数据库初始化脚本
│       └── application*.yml         # 配置文件
│
├── career-planning-frontend/         # Vue 前端应用
│   ├── src/
│   │   ├── views/                   # 页面级组件
│   │   ├── components/              # 公共组件
│   │   ├── api/                     # API 接口封装
│   │   ├── stores/                  # Pinia 状态管理
│   │   ├── router/                  # 路由配置
│   │   ├── utils/                   # 工具函数
│   │   └── assets/                  # 静态资源
│   └── public/                      # 公共静态文件
│
└── career-planning-ai/               # Python AI 服务
    ├── src/
    │   └── ai_service/
    │       ├── agents/              # AI 智能体实现
    │       ├── engine/              # AI 引擎核心
    │       ├── models/              # 数据模型
    │       └── utils/               # 工具函数
    ├── prompts/                     # Prompt 模板
    ├── data/                        # 本地数据存储
    └── config.yaml                  # 配置文件"""

    add_code_block(doc, project_structure)

    doc.add_page_break()

    # ==================== 第二章 技术架构 ====================
    add_heading(doc, "第二章 技术架构", level=1)

    add_heading(doc, "2.1 整体架构设计", level=2)
    add_paragraph(doc,
        "本系统采用前后端分离的微服务架构设计，将业务逻辑与 AI 能力解耦，提升系统的可维护性和可扩展性。")

    add_heading(doc, "2.1.1 系统分层架构", level=3)

    layer_architecture = [
        ("客户端层", "Web 浏览器、移动端 (响应式设计)"),
        ("前端展示层", "Vue 3 + Element Plus + ECharts"),
        ("API 网关层", "Spring Boot Controller + JWT 鉴权"),
        ("业务服务层", "用户服务、简历服务、面试服务、职位服务、支付服务"),
        ("AI 服务层", "对话智能体、简历优化、面试题生成、知识图谱查询"),
        ("数据存储层", "MySQL 数据库、Redis 缓存、RabbitMQ 消息队列"),
        ("基础设施层", "Nginx 反向代理、文件存储、日志系统"),
    ]
    add_table(doc, ["层级", "说明"], layer_architecture, [3.5, 7.5])

    add_heading(doc, "2.1.2 数据流向", level=3)
    add_paragraph(doc,
        "用户请求 → Nginx → 前端应用 → Axios → Spring Boot 后端 → Service 层 → MyBatis Mapper → MySQL/Redis")
    add_paragraph(doc,
        "AI 请求 → Spring Boot → HTTP Client → FastAPI AI 服务 → LLM/API → 返回结果")

    add_heading(doc, "2.2 后端技术栈", level=2)

    add_heading(doc, "2.2.1 核心框架", level=3)

    backend_tech = [
        ("Java", "1.8+", "编程语言基础"),
        ("Spring Boot", "2.7.x", "后端应用框架，简化配置与开发"),
        ("Spring Security", "5.7.x", "安全框架，负责认证与授权"),
        ("MyBatis-Plus", "3.5.x", "ORM 框架增强，简化数据库操作"),
        ("Lombok", "1.18.x", "简化 Java 代码，减少样板代码"),
    ]
    add_table(doc, ["技术组件", "版本", "用途说明"], backend_tech, [2.5, 2, 6.5])

    add_heading(doc, "2.2.2 中间件", level=3)

    middleware_tech = [
        ("MySQL", "8.0.x", "关系型数据库，存储核心业务数据"),
        ("Redis", "7.x", "内存缓存，用于 Session 管理、热点数据缓存"),
        ("RabbitMQ", "3.12.x", "消息队列，用于异步任务和解耦"),
        ("JWT", "0.9.1", "JSON Web Token，无状态认证机制"),
    ]
    add_table(doc, ["中间件", "版本", "用途说明"], middleware_tech, [2.5, 2, 6.5])

    add_heading(doc, "2.2.3 开发与运维工具", level=3)

    devops_tech = [
        ("Maven", "3.8+", "项目构建与依赖管理"),
        ("Git", "最新版本", "版本控制"),
        ("Swagger/SpringDoc", "5.x", "API 文档自动生成"),
        ("Docker", "20.10+", "容器化部署"),
        ("Jenkins", "最新版", "持续集成/持续部署"),
    ]
    add_table(doc, ["工具名称", "版本", "用途说明"], devops_tech, [2.5, 2, 6.5])

    add_heading(doc, "2.3 前端技术栈", level=2)

    frontend_tech = [
        ("Vue.js", "3.4.x", "渐进式前端框架，响应式数据绑定"),
        ("TypeScript", "5.x", "JavaScript 超集，提供类型安全"),
        ("Vite", "5.x", "下一代前端构建工具，极速冷启动"),
        ("Pinia", "2.x", "Vue 官方状态管理库，替代 Vuex"),
        ("Vue Router", "4.x", "官方路由管理器"),
        ("Element Plus", "2.5.x", "基于 Vue 3 的 UI 组件库"),
        ("Axios", "1.6.x", "HTTP 客户端，发起网络请求"),
        ("ECharts", "5.5.x", "百度开源的数据可视化库"),
        ("SCSS", "最新", "CSS 预处理器，增强 CSS 功能"),
    ]
    add_table(doc, ["技术组件", "版本", "用途说明"], frontend_tech, [2.5, 2, 6.5])

    add_heading(doc, "2.4 AI 服务技术栈", level=2)

    ai_tech = [
        ("Python", "3.12", "AI 服务编程语言"),
        ("FastAPI", "0.109.x", "高性能 Python Web 框架，原生支持异步"),
        ("LangChain", "0.2.x", "LLM 应用开发框架，链式调用"),
        ("Pydantic", "2.5.x", "数据验证，使用类型注解"),
        ("Uvicorn", "0.27.x", "ASGI 服务器，运行 FastAPI"),
        ("Loguru", "0.7.x", "改进的日志记录库"),
        ("PaddleOCR", "2.7.x", "百度 OCR 工具，简历解析"),
    ]
    add_table(doc, ["技术组件", "版本", "用途说明"], ai_tech, [2.5, 2, 6.5])

    add_heading(doc, "2.4.1 AI 模型与服务", level=3)

    llm_services = [
        ("通义千问 Qwen", "阿里云 DashScope API", "主要对话模型，职业规划咨询"),
        ("PaddleOCR", "本地部署", "简历 PDF 解析"),
        ("Milvus", "本地/远程", "向量数据库，知识检索增强"),
    ]
    add_table(doc, ["服务/模型", "部署方式", "用途"], llm_services, [3, 4.5, 5.5])

    doc.add_page_break()

    # ==================== 第三章 数据库设计 ====================
    add_heading(doc, "第三章 数据库设计", level=1)

    add_heading(doc, "3.1 数据库概述", level=2)

    db_info = [
        ("数据库类型", "MySQL 8.0"),
        ("数据库名称", "career_backend"),
        ("字符集", "utf8mb4"),
        ("排序规则", "utf8mb4_unicode_ci"),
        ("存储引擎", "InnoDB"),
        ("默认时区", "Asia/Shanghai (+8)"),
    ]
    add_table(doc, ["配置项", "值"], db_info, [4, 7])

    add_heading(doc, "3.1.1 数据库结构概览", level=3)

    table_categories = [
        ("用户相关表", "user, user_profile, feedback"),
        ("简历相关表", "resume, resume_education, resume_work_experience, resume_project, resume_skill"),
        ("面试相关表", "interview, interview_question, interview_review"),
        ("职位相关表", "job, job_skill, user_job_match"),
        ("会员相关表", "member_package, member_subscription, payment_order, points_transaction"),
        ("知识库相关表", "career_path, skill_tree, industry_knowledge"),
        ("系统相关表", "file_upload, system_config, operation_log"),
    ]
    add_table(doc, ["分类", "包含表"], table_categories, [4, 7])

    add_heading(doc, "3.2 核心表设计", level=2)

    add_heading(doc, "3.2.1 用户表 (user)", level=3)
    add_paragraph(doc, "存储系统用户基本信息，包括登录凭证和账户状态。")

    user_fields = [
        ("id", "BIGINT", "主键，自增"),
        ("username", "VARCHAR(50)", "用户名，唯一索引"),
        ("password", "VARCHAR(255)", "密码 (BCrypt 加密)"),
        ("email", "VARCHAR(100)", "邮箱，唯一索引"),
        ("phone", "VARCHAR(20)", "手机号，唯一索引"),
        ("avatar", "VARCHAR(500)", "头像 URL"),
        ("status", "TINYINT", "状态：0-禁用，1-正常"),
        ("last_login_time", "DATETIME", "最后登录时间"),
        ("last_login_ip", "VARCHAR(50)", "最后登录 IP"),
        ("create_time", "DATETIME", "创建时间"),
        ("update_time", "DATETIME", "更新时间"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], user_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.2 用户档案表 (user_profile)", level=3)
    add_paragraph(doc, "扩展用户信息，存储个人背景和职业期望。")

    profile_fields = [
        ("id", "BIGINT", "主键"),
        ("user_id", "BIGINT", "关联用户 ID，唯一索引"),
        ("real_name", "VARCHAR(50)", "真实姓名"),
        ("gender", "TINYINT", "性别：0-未知，1-男，2-女"),
        ("university", "VARCHAR(100)", "院校"),
        ("major", "VARCHAR(100)", "专业"),
        ("education", "TINYINT", "学历：1-大专，2-本科，3-硕士，4-博士"),
        ("graduation_year", "INT", "毕业年份"),
        ("expected_salary", "VARCHAR(50)", "期望薪资"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], profile_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.3 简历表 (resume)", level=3)
    add_paragraph(doc, "存储用户简历，支持多份简历管理。")

    resume_fields = [
        ("id", "BIGINT", "主键"),
        ("user_id", "BIGINT", "用户 ID，索引"),
        ("title", "VARCHAR(100)", "简历标题"),
        ("template_id", "INT", "模板 ID"),
        ("content", "JSON", "简历内容 (结构化 JSON)"),
        ("is_default", "TINYINT", "是否默认简历"),
        ("completeness", "INT", "完整度百分比"),
        ("status", "TINYINT", "状态：0-草稿，1-公开，2-保密"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], resume_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.4 教育经历表 (resume_education)", level=3)

    edu_fields = [
        ("id", "BIGINT", "主键"),
        ("resume_id", "BIGINT", "简历 ID，外键"),
        ("school", "VARCHAR(100)", "学校名称"),
        ("degree", "TINYINT", "学历：1-大专，2-本科，3-硕士，4-博士"),
        ("major", "VARCHAR(100)", "专业"),
        ("start_date", "DATE", "开始时间"),
        ("end_date", "DATE", "结束时间"),
        ("gpa", "DECIMAL(3,2)", "GPA"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], edu_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.5 工作经历表 (resume_work_experience)", level=3)

    work_fields = [
        ("id", "BIGINT", "主键"),
        ("resume_id", "BIGINT", "简历 ID，外键"),
        ("company", "VARCHAR(100)", "公司名称"),
        ("position", "VARCHAR(100)", "职位名称"),
        ("department", "VARCHAR(100)", "部门"),
        ("start_date", "DATE", "开始时间"),
        ("end_date", "DATE", "结束时间"),
        ("description", "TEXT", "工作描述"),
        ("achievements", "TEXT", "工作业绩"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], work_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.6 面试记录表 (interview)", level=3)

    interview_fields = [
        ("id", "BIGINT", "主键"),
        ("user_id", "BIGINT", "用户 ID，索引"),
        ("position", "VARCHAR(100)", "面试职位"),
        ("company", "VARCHAR(100)", "公司名称"),
        ("interview_type", "TINYINT", "类型：1-电话，2-视频，3-现场，4-笔试"),
        ("interview_date", "DATETIME", "面试时间，索引"),
        ("status", "TINYINT", "状态：0-待面试，1-已完成，2-已取消"),
        ("remark", "TEXT", "备注"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], interview_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.7 面试题目表 (interview_question)", level=3)

    question_fields = [
        ("id", "BIGINT", "主键"),
        ("interview_id", "BIGINT", "面试 ID，外键"),
        ("question_type", "TINYINT", "类型：1-技术题，2-行为题，3-场景题"),
        ("question", "TEXT", "题目内容"),
        ("answer", "TEXT", "参考答案"),
        ("difficulty", "TINYINT", "难度：1-简单，2-中等，3-困难"),
        ("is_ai_generated", "TINYINT", "是否 AI 生成"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], question_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.8 职位表 (job)", level=3)

    job_fields = [
        ("id", "BIGINT", "主键"),
        ("title", "VARCHAR(100)", "职位名称，索引"),
        ("company", "VARCHAR(100)", "公司名称，索引"),
        ("industry", "VARCHAR(50)", "行业，索引"),
        ("location", "VARCHAR(100)", "工作地点，索引"),
        ("salary_min", "INT", "最低薪资 (K/月)"),
        ("salary_max", "INT", "最高薪资 (K/月)"),
        ("experience_min", "INT", "最低经验要求 (年)"),
        ("education", "TINYINT", "学历要求"),
        ("description", "TEXT", "职位描述"),
        ("skills_required", "VARCHAR(500)", "所需技能 (JSON)"),
        ("publish_date", "DATE", "发布日期"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], job_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.9 会员套餐表 (member_package)", level=3)

    package_fields = [
        ("id", "BIGINT", "主键"),
        ("name", "VARCHAR(50)", "套餐名称"),
        ("price", "DECIMAL(10,2)", "价格"),
        ("duration_days", "INT", "时长 (天)"),
        ("points_gift", "INT", "赠送积分"),
        ("features", "JSON", "功能权限列表"),
        ("status", "TINYINT", "状态：0-停售，1-在售"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], package_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.10 支付订单表 (payment_order)", level=3)

    order_fields = [
        ("id", "BIGINT", "主键"),
        ("order_no", "VARCHAR(64)", "订单号，唯一索引"),
        ("user_id", "BIGINT", "用户 ID，索引"),
        ("product_type", "TINYINT", "类型：1-会员，2-积分"),
        ("amount", "DECIMAL(10,2)", "支付金额"),
        ("pay_method", "VARCHAR(20)", "支付方式:alipay/wechat"),
        ("transaction_no", "VARCHAR(100)", "第三方交易号"),
        ("status", "TINYINT", "状态:0-待支付，1-已支付，2-已取消"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], order_fields, [3, 2.5, 5.5])

    add_heading(doc, "3.2.11 积分交易表 (points_transaction)", level=3)

    points_fields = [
        ("id", "BIGINT", "主键"),
        ("user_id", "BIGINT", "用户 ID，索引"),
        ("transaction_type", "TINYINT", "类型：1-收入，2-支出"),
        ("points", "INT", "积分数"),
        ("balance", "INT", "交易后余额"),
        ("source", "VARCHAR(50)", "来源:signin/interview/member/payment"),
        ("description", "VARCHAR(255)", "描述"),
    ]
    add_table(doc, ["字段名", "数据类型", "说明"], points_fields, [3, 2.5, 5.5])

    doc.add_page_break()

    # ==================== 第四章 API 接口文档 ====================
    add_heading(doc, "第四章 API 接口文档", level=1)

    add_heading(doc, "4.1 通用说明", level=2)

    api_common = [
        ("基础 URL", "http://localhost:8080/api"),
        ("AI 服务 URL", "http://localhost:9000"),
        ("认证方式", "JWT Token (Bearer Token)"),
        ("请求格式", "application/json"),
        ("文件上传", "multipart/form-data"),
        ("字符编码", "UTF-8"),
    ]
    add_table(doc, ["配置项", "值"], api_common, [4, 7])

    add_heading(doc, "4.1.1 统一响应格式", level=3)

    response_example = """{
  "code": 200,
  "state": true,
  "msg": "操作成功",
  "data": { ... },
  "timestamp": 1712995200000
}

字段说明:
- code: 状态码 (200-成功，其他 - 错误)
- state: 布尔值，表示操作是否成功
- msg: 提示信息
- data: 业务数据
- timestamp: 时间戳 (毫秒)"""

    add_code_block(doc, response_example)

    add_heading(doc, "4.1.2 错误码说明", level=3)

    error_codes = [
        ("200", "成功"),
        ("400", "请求参数错误"),
        ("401", "未认证/Token 过期"),
        ("403", "权限不足"),
        ("404", "资源不存在"),
        ("500", "服务器内部错误"),
        ("1001", "用户不存在"),
        ("1002", "密码错误"),
        ("1003", "验证码错误或过期"),
        ("2001", "简历解析失败"),
        ("3001", "AI 服务不可用"),
    ]
    add_table(doc, ["错误码", "说明"], error_codes, [4, 7])

    add_heading(doc, "4.2 用户模块", level=2)

    user_apis = [
        ("用户登录", "POST", "/user/login", "email, password, userType"),
        ("用户注册", "POST", "/user/register", "email, password, code, userType"),
        ("忘记密码", "POST", "/user/forget-password", "email, password, code"),
        ("获取用户信息", "GET", "/user/info", "Header: Authorization"),
        ("更新用户信息", "PUT", "/user/edit", "nickname, gender, bio"),
        ("上传头像", "POST", "/user/avatar", "multipart/form-data: file"),
        ("刷新 Token", "POST", "/user/refreshToken", "refreshToken"),
        ("发送验证码", "GET", "/user/sendCode", "email"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], user_apis, [3, 1.5, 2.5, 6])

    add_heading(doc, "4.3 简历模块", level=2)

    resume_apis = [
        ("简历解析", "POST", "/api/parse/resume", "multipart/form-data: file(PDF/Word)"),
        ("获取简历列表", "GET", "/api/resumes", "无"),
        ("获取简历详情", "GET", "/api/resumes/{id}", "简历 ID"),
        ("创建/更新简历", "POST/PUT", "/api/resumes", "简历内容 JSON"),
        ("删除简历", "DELETE", "/api/resumes/{id}", "简历 ID"),
        ("简历导出", "POST", "/api/resume/export", "resumeId, format, templateId"),
        ("设置默认简历", "PUT", "/api/resumes/{id}/default", "简历 ID"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], resume_apis, [2.5, 1.5, 2.5, 7])

    add_heading(doc, "4.4 面试管理模块", level=2)

    interview_apis = [
        ("获取面试列表", "GET", "/api/interviews", "page, size, status"),
        ("添加面试记录", "POST", "/api/interviews", "company, position, interviewTime"),
        ("获取面试详情", "GET", "/api/interviews/{id}", "面试 ID"),
        ("更新面试记录", "PUT", "/api/interviews/{id}", "修改字段"),
        ("删除面试记录", "DELETE", "/api/interviews/{id}", "面试 ID"),
        ("面试回顾", "POST", "/api/interviews/review", "interviewId, content, questions, rating"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], interview_apis, [2.5, 1.5, 2.5, 7])

    add_heading(doc, "4.5 AI 智能助手模块", level=2)

    ai_apis = [
        ("AI 对话", "POST", "/api/chat/send", "message, userId"),
        ("获取对话历史", "GET", "/api/chat/history/{userId}", "用户 ID"),
        ("清空对话历史", "DELETE", "/api/chat/history/{userId}", "用户 ID"),
        ("面试题生成", "POST", "/api/question/generate", "jobTitle, difficulty, count"),
        ("代码能力评估", "POST", "/api/code/ability", "language, code"),
        ("简历优化建议", "POST", "/api/resume/optimize", "resumeId"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], ai_apis, [2.5, 1.5, 2.5, 7])

    add_heading(doc, "4.6 职位匹配模块", level=2)

    job_apis = [
        ("职位搜索", "GET", "/api/search/positions", "keyword, city, salary, page, size"),
        ("获取职位详情", "GET", "/api/jobs/{id}", "职位 ID"),
        ("推荐职位", "GET", "/api/jobs/recommend", "userId, limit"),
        ("职位匹配分析", "POST", "/api/match/job", "resumeId"),
        ("收藏职位", "POST", "/api/jobs/favorite", "jobId"),
        ("取消收藏", "DELETE", "/api/jobs/favorite/{jobId}", "职位 ID"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], job_apis, [2.5, 1.5, 2.5, 7])

    add_heading(doc, "4.7 会员与支付模块", level=2)

    pay_apis = [
        ("获取会员套餐列表", "GET", "/api/packages", "无"),
        ("购买套餐", "POST", "/api/package/buy", "packageId, payMethod"),
        ("获取订单详情", "GET", "/api/orders/{orderId}", "订单 ID"),
        ("支付宝回调", "POST", "/api/pay/callback", "支付宝异步通知"),
        ("积分充值", "POST", "/api/points/recharge", "points, payMethod"),
        ("获取积分明细", "GET", "/api/points/transactions", "page, size"),
        ("会员信息查询", "GET", "/api/member/info", "无"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], pay_apis, [2.5, 1.5, 2.5, 7])

    add_heading(doc, "4.8 用户反馈模块", level=2)

    feedback_apis = [
        ("提交反馈", "POST", "/api/feedback", "feedbackType, title, content, images"),
        ("获取反馈列表", "GET", "/api/feedback", "page, size"),
        ("反馈详情", "GET", "/api/feedback/{id}", "反馈 ID"),
        ("回复反馈", "PUT", "/api/feedback/{id}/reply", "replyContent"),
    ]
    add_table(doc, ["接口名称", "方法", "路径", "主要参数"], feedback_apis, [2.5, 1.5, 2.5, 7])

    doc.add_page_break()

    # ==================== 第五章 部署指南 ====================
    add_heading(doc, "第五章 部署指南", level=1)

    add_heading(doc, "5.1 环境要求", level=2)

    env_requirements = [
        ("JDK", "1.8", "11/17", "后端运行环境"),
        ("Node.js", "16.x", "20.x LTS", "前端构建环境"),
        ("Python", "3.10", "3.12.x", "AI 服务运行环境"),
        ("MySQL", "8.0", "8.0.x", "主数据库"),
        ("Redis", "6.x", "7.x", "缓存/会话存储"),
        ("RabbitMQ", "3.10", "3.12.x", "消息队列"),
        ("Nginx", "1.20", "1.24", "Web 服务器"),
    ]
    add_table(doc, ["组件", "最低版本", "推荐版本", "说明"], env_requirements, [2, 2, 2, 5])

    add_heading(doc, "5.1.2 服务器配置", level=3)

    server_specs = [
        ("CPU", "2 核", "4 核及以上"),
        ("内存", "4GB", "8GB 及以上"),
        ("硬盘", "50GB", "100GB 及以上 SSD"),
        ("带宽", "2Mbps", "5Mbps 及以上"),
    ]
    add_table(doc, ["配置项", "最低配置", "推荐配置"], server_specs, [3, 3, 4])

    add_heading(doc, "5.2 部署架构", level=2)

    add_paragraph(doc,
        "系统采用 Nginx 作为反向代理服务器，前端静态资源由 Nginx 直接服务，"
        "API 请求通过 Nginx 转发至 Spring Boot 后端服务 (端口 8080)，"
        "后端服务与 AI 服务 (端口 9000) 通过 HTTP 通信。"
        "基础设施层包括 MySQL(3306)、Redis(6379)、RabbitMQ(5672) 等服务。")

    add_heading(doc, "5.3 后端部署", level=2)

    add_heading(doc, "5.3.1 构建项目", level=3)

    build_cmds = """# 进入后端目录
cd career-planning-backend

# Maven 构建
mvn clean package -DskipTests

# 构建产物位置
ls target/*.jar"""

    add_code_block(doc, build_cmds)

    add_heading(doc, "5.3.2 配置文件", level=3)
    add_paragraph(doc, "编辑 application-prod.yaml 配置生产环境参数:")

    config_yaml = """spring:
  datasource:
    url: jdbc:mysql://localhost:3306/career_backend?useUnicode=true&characterEncoding=utf8
    username: career
    password: your_password
  redis:
    host: localhost
    port: 6379
    password: your_redis_password

jwt:
  secret: your_jwt_secret_key_here
  expiration: 86400000

ai:
  service:
    url: http://localhost:9000
    timeout: 30000"""

    add_code_block(doc, config_yaml)

    add_heading(doc, "5.3.3 Systemd 服务", level=3)

    systemd_config = """[Unit]
Description=Career Planning Backend Service
After=network.target mysql.service redis.service

[Service]
Type=simple
User=root
WorkingDirectory=/opt/career-backend
ExecStart=/usr/bin/java -jar /opt/career-backend/career-planning-backend.jar
Restart=on-failure
RestartSec=10

[Install]
WantedBy=multi-user.target"""

    add_code_block(doc, systemd_config)

    add_paragraph(doc, "启动服务:")
    start_cmds = """sudo systemctl daemon-reload
sudo systemctl enable career-backend
sudo systemctl start career-backend
sudo systemctl status career-backend"""
    add_code_block(doc, start_cmds)

    add_heading(doc, "5.4 前端部署", level=2)

    frontend_build = """# 安装依赖
npm install

# 配置环境变量
cp .env.production.example .env.production
# 编辑.env.production 设置 VITE_API_BASE_URL

# 构建生产版本
npm run build

# 部署到 Nginx
sudo cp -r dist/* /var/www/career-frontend/"""

    add_code_block(doc, frontend_build)

    add_heading(doc, "5.4.1 Nginx 配置", level=3)

    nginx_conf = """server {
    listen 80;
    server_name www.your-domain.com;

    root /var/www/career-frontend/dist;
    index index.html;

    # SPA 路由支持
    location / {
        try_files $uri $uri/ /index.html;
    }

    # API 代理
    location /api/ {
        proxy_pass http://localhost:8080;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
    }
}"""

    add_code_block(doc, nginx_conf)

    add_heading(doc, "5.5 AI 服务部署", level=2)

    add_heading(doc, "5.5.1 Docker 部署 (推荐)", level=3)

    docker_deploy = """# 构建镜像
docker build -t career-ai:latest .

# 运行容器
docker run -d \\
  --name career-ai \\
  -p 9000:9000 \\
  -v /opt/career-ai/config.yaml:/app/config.yaml \\
  career-ai:latest

# 查看日志
docker logs -f career-ai"""

    add_code_block(doc, docker_deploy)

    add_heading(doc, "5.6 监控与日志", level=2)

    log_paths = [
        ("/var/log/career-backend/", "后端应用日志"),
        ("/var/log/career-ai/", "AI 服务日志"),
        ("/var/log/nginx/", "Nginx 访问/错误日志"),
    ]
    add_table(doc, ["日志路径", "说明"], log_paths, [5, 5])

    add_heading(doc, "5.6.1 健康检查", level=3)

    health_cmds = """# 后端健康检查
curl http://localhost:8080/actuator/health

# AI 服务健康检查
curl http://localhost:9000/

# Redis 连接测试
redis-cli ping

# MySQL 连接测试
mysql -u career -p -e "SELECT 1" """

    add_code_block(doc, health_cmds)

    add_heading(doc, "5.7 数据库备份", level=2)

    backup_script = """#!/bin/bash
BACKUP_DIR=/opt/backup/mysql
DATE=$(date +%Y%m%d_%H%M%S)

mysqldump -ucareer -ppassword career_backend | gzip > \\
  $BACKUP_DIR/career_${DATE}.sql.gz

# 保留最近 7 天的备份
find $BACKUP_DIR -name "*.sql.gz" -mtime +7 -delete"""

    add_code_block(doc, backup_script)

    add_paragraph(doc, "添加到 crontab 每日凌晨 2 点自动备份:")
    cron_cmd = """0 2 * * * /opt/backup/backup-mysql.sh"""
    add_code_block(doc, cron_cmd)

    doc.add_page_break()

    # ==================== 附录 ====================
    add_heading(doc, "附录 A: 快速开始", level=1)

    add_heading(doc, "A.1 本地开发环境搭建", level=2)

    quick_start = """# 1. 克隆项目
git clone https://github.com/your-org/CareerAgent.git
cd CareerAgent

# 2. 初始化数据库
mysql -u root -p < career-planning-backend/src/main/resources/sql/schema.sql
mysql -u root -p career_backend < career-planning-backend/src/main/resources/sql/data.sql

# 3. 启动后端
cd career-planning-backend
mvn spring-boot:run

# 4. 启动前端
cd career-planning-frontend
npm install
npm run dev

# 5. 启动 AI 服务 (可选)
cd career-planning-ai
poetry install
poetry run uvicorn main:app --reload --port 9000

# 访问应用
# 前端：http://localhost:5173
# 后端 API: http://localhost:8080/api
# AI 服务：http://localhost:9000"""

    add_code_block(doc, quick_start)

    add_heading(doc, "附录 B: API 调试示例", level=1)

    add_heading(doc, "B.1 用户登录", level=2)

    login_example = """curl -X POST http://localhost:8080/api/user/login \\
  -H "Content-Type: application/json" \\
  -d '{
    "email": "test@example.com",
    "password": "your_password",
    "userType": "student"
  }'

# 响应示例
{
  "code": 200,
  "state": true,
  "msg": "登录成功",
  "data": {
    "token": "eyJhbGciOiJIUzI1NiIs...",
    "user": {
      "id": 1,
      "username": "testuser",
      "email": "test@example.com"
    }
  }
}"""

    add_code_block(doc, login_example)

    add_heading(doc, "B.2 AI 对话", level=2)

    chat_example = """curl -X POST http://localhost:8080/api/chat/send \\
  -H "Content-Type: application/json" \\
  -H "Authorization: Bearer YOUR_TOKEN" \\
  -d '{
    "message": "我是一名计算机科学专业的应届毕业生，想从事软件开发工作，请给我一些建议。",
    "userId": 1
  }'

# 响应示例
{
  "code": 200,
  "state": true,
  "msg": "success",
  "data": {
    "reply": "作为一名计算机专业的应届毕业生，我给你以下几点建议:\n\n1. 夯实基础...\n2. 项目经验...\n3. ..."
  }
}"""

    add_code_block(doc, chat_example)

    # 保存文档
    output_path = "md/职引未来_技术文档_v5.0.docx"
    doc.save(output_path)
    print(f"文档已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
