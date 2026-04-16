#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成职引未来技术文档Word文件 - 增强版
按照参考文档格式进行美化，添加表格和更专业的排版
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 设置中文字体
def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    if level == 1:
        # 一级标题 - 黑体三号加粗居中
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 16, bold=True)
        p.space_after = Pt(18)
        p.space_before = Pt(12)
    elif level == 2:
        # 二级标题 - 黑体四号加粗
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 14, bold=True)
        p.space_after = Pt(12)
        p.space_before = Pt(12)
    elif level == 3:
        # 三级标题 - 黑体小四加粗
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 12, bold=True)
        p.space_after = Pt(6)
        p.space_before = Pt(6)
    else:
        # 四级标题 - 宋体小四加粗
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_chinese_font(run, '宋体', 12, bold=True)
        p.space_after = Pt(6)
    return p

def add_paragraph_custom(doc, text, indent=True, bold=False, font_name='宋体', font_size=12):
    """添加自定义段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    run = p.add_run(text)
    set_chinese_font(run, font_name, font_size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.space_after = Pt(6)
    return p

def add_table_custom(doc, headers, rows, col_widths=None):
    """添加自定义表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_chinese_font(run, '黑体', 11, bold=True)

    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for para in row_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_chinese_font(run, '宋体', 10.5)

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    doc.add_paragraph()  # 表格后加空行
    return table

def create_doc_from_content():
    """根据MD文档内容生成Word文档"""

    # 创建文档
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ==================== 封面 ====================
    add_heading_custom(doc, "职引未来", level=1)
    add_heading_custom(doc, "基于AI的大学生职业规划智能体系统", level=1)
    add_heading_custom(doc, "技术文档", level=1)

    # 空行
    for _ in range(8):
        doc.add_paragraph()

    # 文档信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run("\n\n版本：v2.0\n\n")
    set_chinese_font(run, '宋体', 12)
    run = info_para.add_run("日期：2026年4月\n")
    set_chinese_font(run, '宋体', 12)

    # 分页
    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading_custom(doc, "目录", level=1)

    toc_items = [
        ("一、项目概述", "1"),
        ("二、技术架构", "2"),
        ("三、数据库设计", "3"),
        ("四、API接口文档", "4"),
        ("五、部署指南", "5")
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(item)
        set_chinese_font(run, '宋体', 12)
        # 添加页码（右对齐）
        tab_run = p.add_run("\t" + page)
        set_chinese_font(tab_run, '宋体', 12)

    doc.add_page_break()

    # ==================== 一、项目概述 ====================
    add_heading_custom(doc, "一、项目概述", level=1)

    add_heading_custom(doc, "1.1 系统简介", level=2)
    add_paragraph_custom(doc,
        "「职引未来」是一个基于AI的大学生职业规划智能体系统，旨在帮助大学生进行科学的职业规划与就业准备。"
        "系统整合了职业测评、简历优化、面试准备、职业知识库、AI智能问答等多维度功能，为用户提供全方位的职业发展支持。")

    add_heading_custom(doc, "1.2 核心功能模块", level=2)

    modules = [
        ("用户管理模块", "用户注册与登录、用户信息管理、会员体系管理、积分系统管理"),
        ("职业测评模块", "职业能力评估、技能雷达图展示、个人优势分析"),
        ("简历管理模块", "简历在线编辑、简历模板选择、简历解析与智能填充、简历导出功能"),
        ("AI智能助手模块", "AI职业规划问答、面试题智能生成、职业发展建议、用户画像分析"),
        ("面试管理模块", "面试题库管理、面试日历、面试回顾与反馈"),
        ("职位匹配模块", "职位搜索、智能岗位推荐、用户与岗位匹配分析"),
        ("知识库模块", "职业知识图谱、行业知识浏览、发展路径导航")
    ]

    for name, desc in modules:
        add_paragraph_custom(doc, f"（1）{name}：{desc}")

    add_heading_custom(doc, "1.3 技术栈概览", level=2)

    # 技术栈表格
    tech_data = [
        ("后端框架", "Spring Boot 2.7.x"),
        ("前端框架", "Vue 3 + TypeScript"),
        ("AI 服务", "FastAPI + Python"),
        ("数据库", "MySQL 8.0"),
        ("ORM 框架", "MyBatis-Plus"),
        ("缓存", "Redis"),
        ("消息队列", "RabbitMQ"),
        ("API 文档", "Swagger"),
        ("AI 模型", "通义千问 (Qwen)")
    ]
    add_table_custom(doc, ["层级", "技术选型"], tech_data, [3, 8])

    add_heading_custom(doc, "1.4 项目结构", level=2)
    add_paragraph_custom(doc, "CareerAgent/")
    add_paragraph_custom(doc, "├── career-planning-backend/     # Java 后端服务", indent=False)
    add_paragraph_custom(doc, "│   ├── src/main/java/", indent=False)
    add_paragraph_custom(doc, "│   │   ├── config/             # 配置类", indent=False)
    add_paragraph_custom(doc, "│   │   ├── controller/         # 控制器层", indent=False)
    add_paragraph_custom(doc, "│   │   ├── domain/             # 实体与DTO", indent=False)
    add_paragraph_custom(doc, "│   │   ├── service/            # 服务层", indent=False)
    add_paragraph_custom(doc, "│   │   ├── mapper/             # 数据访问层", indent=False)
    add_paragraph_custom(doc, "│   │   └── utils/              # 工具类", indent=False)
    add_paragraph_custom(doc, "│   └── src/main/resources/     # 配置文件", indent=False)
    add_paragraph_custom(doc, "├── career-planning-frontend/    # Vue 前端应用", indent=False)
    add_paragraph_custom(doc, "│   ├── src/views/              # 页面组件", indent=False)
    add_paragraph_custom(doc, "│   ├── components/             # 公共组件", indent=False)
    add_paragraph_custom(doc, "│   ├── api/                    # API 接口封装", indent=False)
    add_paragraph_custom(doc, "│   ├── stores/                 # Pinia 状态管理", indent=False)
    add_paragraph_custom(doc, "│   └── router/                 # 路由配置", indent=False)
    add_paragraph_custom(doc, "└── career-planning-ai/          # Python AI 服务", indent=False)
    add_paragraph_custom(doc, "    ├── ai_service/", indent=False)
    add_paragraph_custom(doc, "    │   ├── agents/             # AI 智能体", indent=False)
    add_paragraph_custom(doc, "    │   ├── engine/             # AI 引擎", indent=False)
    add_paragraph_custom(doc, "    │   └── utils/              # 工具类", indent=False)
    add_paragraph_custom(doc, "    └── main.py                 # 服务入口", indent=False)

    doc.add_page_break()

    # ==================== 二、技术架构 ====================
    add_heading_custom(doc, "二、技术架构", level=1)

    add_heading_custom(doc, "2.1 整体架构设计", level=2)
    add_paragraph_custom(doc,
        "本系统采用前后端分离的微服务架构设计，后端基于Spring Boot框架，AI服务基于FastAPI + Python构建，通过HTTP接口进行通信。")

    add_paragraph_custom(doc, "系统架构分层如下：")
    add_paragraph_custom(doc, "（1）客户端层：Web浏览器、移动端（响应式）")
    add_paragraph_custom(doc, "（2）前端应用层：Vue 3 + TypeScript + Vite + Pinia + Vue Router")
    add_paragraph_custom(doc, "（3）API网关层：Spring Boot 2.7.x + Spring Security + JWT")
    add_paragraph_custom(doc, "（4）业务服务层：用户管理、简历管理、面试管理、职位匹配、支付积分服务")
    add_paragraph_custom(doc, "（5）AI服务层：FastAPI + LangChain + LLM模型调用 + 知识图谱引擎")
    add_paragraph_custom(doc, "（6）基础设施层：MySQL、Redis、RabbitMQ、文件存储")

    add_heading_custom(doc, "2.2 技术栈详情", level=2)

    add_heading_custom(doc, "2.2.1 后端技术栈", level=3)
    backend_data = [
        ("组件", "版本", "用途"),
        ("Java", "1.8+", "编程语言"),
        ("Spring Boot", "2.7.x", "后端框架"),
        ("MyBatis-Plus", "3.5.x", "ORM框架"),
        ("MySQL", "8.0.x", "主数据库"),
        ("Redis", "7.x", "缓存/会话"),
        ("RabbitMQ", "3.12.x", "消息队列"),
        ("JWT", "0.9.1", "Token认证"),
        ("Lombok", "1.18.x", "代码简化"),
        ("Swagger", "3.0.x", "API文档")
    ]
    add_table_custom(doc, ["组件", "版本", "用途"], backend_data[1:], [3, 3, 6])

    add_heading_custom(doc, "2.2.2 前端技术栈", level=3)
    frontend_data = [
        ("Vue", "3.4.x", "前端框架"),
        ("TypeScript", "5.x", "类型安全"),
        ("Vite", "5.x", "构建工具"),
        ("Pinia", "2.x", "状态管理"),
        ("Vue Router", "4.x", "路由管理"),
        ("Element Plus", "2.5.x", "UI组件库"),
        ("Axios", "1.6.x", "HTTP客户端"),
        ("ECharts", "5.5.x", "图表展示"),
        ("SCSS", "1.69.x", "CSS预处理器")
    ]
    add_table_custom(doc, ["组件", "版本", "用途"], frontend_data, [3, 3, 6])

    add_heading_custom(doc, "2.2.3 AI服务技术栈", level=3)
    ai_data = [
        ("Python", "3.12", "编程语言"),
        ("FastAPI", "0.109.x", "Web框架"),
        ("LangChain", "0.2.x", "LLM应用框架"),
        ("LangGraph", "0.1.x", "工作流引擎"),
        ("Pydantic", "2.5.x", "数据验证"),
        ("Uvicorn", "0.27.x", "ASGI服务器"),
        ("Loguru", "0.7.x", "日志框架")
    ]
    add_table_custom(doc, ["组件", "版本", "用途"], ai_data, [3, 3, 6])

    add_heading_custom(doc, "2.3 核心模块设计", level=2)

    add_heading_custom(doc, "2.3.1 用户认证模块", level=3)
    add_paragraph_custom(doc, "安全特性：")
    add_paragraph_custom(doc, "（1）JWT Token认证（RS256算法）")
    add_paragraph_custom(doc, "（2）Token自动续期机制")
    add_paragraph_custom(doc, "（3）密码加密存储（BCrypt）")
    add_paragraph_custom(doc, "（4）接口访问频率限制")

    add_heading_custom(doc, "2.3.2 AI智能体架构", level=3)
    add_paragraph_custom(doc, "智能体类型：")
    add_paragraph_custom(doc, "（1）CareerAnalystAgent - 职业分析师")
    add_paragraph_custom(doc, "（2）GrowthPlanAgent - 成长规划师")
    add_paragraph_custom(doc, "（3）ConversationAgent - 对话助手")
    add_paragraph_custom(doc, "（4）TestQuestionAgent - 面试题生成器")
    add_paragraph_custom(doc, "（5）ExplainAgent - 解释说明器")
    add_paragraph_custom(doc, "（6）UserJobMatchAnalyzer - 岗位匹配分析器")

    add_heading_custom(doc, "2.3.3 消息队列设计", level=3)
    add_paragraph_custom(doc, "RabbitMQ消息队列包括：订单队列(order_queue)、日志队列(log_queue)、通知队列(notify_queue)")

    add_heading_custom(doc, "2.3.4 缓存策略", level=3)
    cache_data = [
        ("缓存数据", "缓存策略", "过期时间"),
        ("用户会话", "Redis String", "30分钟"),
        ("热点职位", "Redis Hash", "1小时"),
        ("面试题库", "Redis List", "24小时"),
        ("AI对话上下文", "Redis", "会话结束")
    ]
    add_table_custom(doc, ["缓存数据", "缓存策略", "过期时间"], cache_data[1:], [3, 3, 3])

    doc.add_page_break()

    # ==================== 三、数据库设计 ====================
    add_heading_custom(doc, "三、数据库设计", level=1)

    add_heading_custom(doc, "3.1 数据库概述", level=2)

    db_info = [
        ("项目", "说明"),
        ("数据库类型", "MySQL 8.0"),
        ("数据库名称", "career_backend"),
        ("字符集", "utf8mb4"),
        ("排序规则", "utf8mb4_unicode_ci"),
        ("存储引擎", "InnoDB")
    ]
    add_table_custom(doc, ["项目", "说明"], db_info[1:], [3, 8])

    add_heading_custom(doc, "3.2 数据库结构概览", level=2)
    add_paragraph_custom(doc, "career_backend/")
    add_paragraph_custom(doc, "├── 用户相关表：user（用户表）、user_profile（用户档案表）、feedback（用户反馈表）", indent=False)
    add_paragraph_custom(doc, "├── 简历相关表：resume（简历表）、resume_education（教育经历表）、resume_work_experience（工作经历表）、resume_project（项目经历表）、resume_skill（技能表）", indent=False)
    add_paragraph_custom(doc, "├── 面试相关表：interview（面试记录表）、interview_question（面试题目表）、interview_review（面试回顾表）", indent=False)
    add_paragraph_custom(doc, "├── 职位相关表：job（职位表）、job_skill（职位技能要求表）、user_job_match（用户-职位匹配表）", indent=False)
    add_paragraph_custom(doc, "├── 会员相关表：member_package（会员套餐表）、member_subscription（会员订阅表）、payment_order（支付订单表）、points_transaction（积分交易表）", indent=False)
    add_paragraph_custom(doc, "├── 知识库相关表：career_path（职业发展路径表）、skill_tree（技能树表）、industry_knowledge（行业知识表）", indent=False)
    add_paragraph_custom(doc, "└── 系统相关表：file_upload（文件上传记录表）、system_config（系统配置表）、operation_log（操作日志表）", indent=False)

    add_heading_custom(doc, "3.3 核心表设计", level=2)

    tables_info = [
        ("user", "用户表", "id, username, password, email, phone, avatar, status"),
        ("user_profile", "用户档案表", "user_id, real_name, gender, education, university"),
        ("resume", "简历表", "user_id, title, content, template_id, status"),
        ("resume_education", "教育经历表", "resume_id, school, degree, major, start_date"),
        ("resume_work_experience", "工作经历表", "resume_id, company, position, start_date"),
        ("resume_project", "项目经历表", "resume_id, project_name, role, technology_stack"),
        ("resume_skill", "技能表", "resume_id, skill_name, proficiency_level"),
        ("interview", "面试记录表", "user_id, position, company, interview_date, status"),
        ("interview_question", "面试题表", "interview_id, question, answer, difficulty"),
        ("job", "职位表", "title, company, location, salary_min, salary_max"),
        ("member_package", "会员套餐表", "name, price, duration_days, points_gift"),
        ("payment_order", "支付订单表", "order_no, user_id, amount, status"),
        ("points_transaction", "积分交易表", "user_id, transaction_type, points, balance"),
        ("feedback", "用户反馈表", "user_id, feedback_type, title, content, status")
    ]

    for i, (table_name, desc, fields) in enumerate(tables_info, 1):
        add_heading_custom(doc, f"3.3.{i} {table_name}表（{desc}）", level=3)
        add_paragraph_custom(doc, f"核心字段：{fields}")

    add_heading_custom(doc, "3.4 索引优化建议", level=2)
    index_data = [
        ("表名", "索引字段", "查询场景"),
        ("user", "phone, email", "登录查询"),
        ("resume", "user_id", "用户简历列表"),
        ("interview", "user_id, interview_date", "面试日历"),
        ("job", "title, location, industry", "职位搜索"),
        ("payment_order", "order_no, status", "订单查询"),
        ("points_transaction", "user_id, create_time", "积分记录")
    ]
    add_table_custom(doc, ["表名", "索引字段", "查询场景"], index_data[1:], [3, 3, 6])

    add_heading_custom(doc, "3.5 分区建议", level=2)
    add_paragraph_custom(doc, "（1）operation_log：按create_time月分区")
    add_paragraph_custom(doc, "（2）payment_order：按create_time月分区")
    add_paragraph_custom(doc, "（3）points_transaction：按user_id + create_time分区")

    doc.add_page_break()

    # ==================== 四、API接口文档 ====================
    add_heading_custom(doc, "四、API接口文档", level=1)

    add_heading_custom(doc, "4.1 通用说明", level=2)

    api_info = [
        ("项目", "说明"),
        ("基础URL", "http://localhost:8080/api"),
        ("AI服务URL", "http://localhost:8000"),
        ("认证方式", "JWT Token (Bearer Token)"),
        ("请求格式", "JSON"),
        ("响应格式", "code, state, msg, data")
    ]
    add_table_custom(doc, ["项目", "说明"], api_info[1:], [3, 8])

    add_heading_custom(doc, "4.2 用户模块", level=2)

    user_apis = [
        ("用户登录", "POST /user/login", "email, password, userType"),
        ("用户注册", "POST /user/register", "email, password, code, userType"),
        ("忘记密码", "POST /user/forget-password", "email, password, code"),
        ("获取用户信息", "GET /user/info", "Authorization: Bearer <token>"),
        ("更新用户信息", "PUT /user/edit", "nickname, gender, bio"),
        ("上传头像", "POST /user/avatar", "multipart/form-data"),
        ("刷新Token", "POST /user/refreshToken", "refreshToken")
    ]
    add_table_custom(doc, ["接口名称", "请求路径", "主要参数"], user_apis, [3, 4, 6])

    add_heading_custom(doc, "4.3 简历模块", level=2)
    add_paragraph_custom(doc, "（1）简历解析：POST /api/parse/resume（Content-Type: multipart/form-data）")
    add_paragraph_custom(doc, "（2）简历导出：POST /api/resume/export（resumeId, format, templateId）")

    add_heading_custom(doc, "4.4 职位匹配模块", level=2)
    add_paragraph_custom(doc, "（1）职位匹配：POST /api/match/job（基于简历进行智能职位匹配）")
    add_paragraph_custom(doc, "（2）职位搜索：GET /api/search/positions（keyword, city, salary, page, size）")

    add_heading_custom(doc, "4.5 AI智能助手模块", level=2)
    ai_apis = [
        ("AI对话", "POST /api/chat/send", "message, userId"),
        ("获取对话历史", "GET /api/chat/history/{userId}", "userId"),
        ("面试题生成", "POST /api/question/generate", "jobTitle, difficulty, count"),
        ("代码能力评估", "POST /api/code/ability", "language, code")
    ]
    add_table_custom(doc, ["接口名称", "请求路径", "主要参数"], ai_apis, [3, 4, 6])

    add_heading_custom(doc, "4.6 会员与支付模块", level=2)
    pay_apis = [
        ("购买套餐", "POST /api/package/buy", "packageId, payMethod"),
        ("支付回调", "POST /api/pay/callback", "支付宝异步通知"),
        ("积分充值", "POST /api/points/recharge", "points, payMethod"),
        ("获取积分明细", "GET /api/points/transactions", "无"),
        ("会员信息查询", "GET /api/member/info", "无")
    ]
    add_table_custom(doc, ["接口名称", "请求路径", "主要参数"], pay_apis, [3, 4, 6])

    add_heading_custom(doc, "4.7 面试管理模块", level=2)
    interview_apis = [
        ("获取面试列表", "GET /api/interviews", "无"),
        ("添加面试记录", "POST /api/interviews", "company, position, interviewTime"),
        ("面试回顾", "POST /api/interviews/review", "interviewId, content, questions, rating")
    ]
    add_table_custom(doc, ["接口名称", "请求路径", "主要参数"], interview_apis, [3, 4, 6])

    add_heading_custom(doc, "4.8 错误码说明", level=2)
    error_data = [
        ("错误码", "说明"),
        ("200", "成功"),
        ("400", "请求参数错误"),
        ("401", "未授权/Token过期"),
        ("403", "禁止访问"),
        ("404", "资源不存在"),
        ("500", "服务器内部错误")
    ]
    add_table_custom(doc, ["错误码", "说明"], error_data[1:], [3, 8])

    doc.add_page_break()

    # ==================== 五、部署指南 ====================
    add_heading_custom(doc, "五、部署指南", level=1)

    add_heading_custom(doc, "5.1 环境要求", level=2)

    add_heading_custom(doc, "5.1.1 基础环境", level=3)
    env_data = [
        ("组件", "最低版本", "推荐版本", "说明"),
        ("JDK", "1.8", "11/17", "后端运行环境"),
        ("Node.js", "16.x", "20.x LTS", "前端构建环境"),
        ("Python", "3.10", "3.12.x", "AI服务运行环境"),
        ("MySQL", "8.0", "8.0.x", "主数据库"),
        ("Redis", "6.x", "7.x", "缓存/会话存储"),
        ("RabbitMQ", "3.10", "3.12.x", "消息队列"),
        ("Nginx", "1.20", "1.24", "Web服务器/反向代理")
    ]
    add_table_custom(doc, ["组件", "最低版本", "推荐版本", "说明"], env_data[1:], [2.5, 2.5, 3, 5])

    add_heading_custom(doc, "5.1.2 服务器配置", level=3)
    server_data = [
        ("配置项", "最低配置", "推荐配置"),
        ("CPU", "2核", "4核及以上"),
        ("内存", "4GB", "8GB及以上"),
        ("硬盘", "50GB", "100GB及以上 SSD"),
        ("带宽", "1Mbps", "5Mbps及以上")
    ]
    add_table_custom(doc, ["配置项", "最低配置", "推荐配置"], server_data[1:], [3, 4, 5])

    add_heading_custom(doc, "5.2 部署架构", level=2)
    add_paragraph_custom(doc,
        "系统采用Nginx作为反向代理，前端静态资源由Nginx直接服务，"
        "API请求转发至Spring Boot后端服务（端口8080），"
        "后端服务与AI服务（端口9000）通过HTTP通信。"
        "基础设施包括MySQL（3306）、Redis（6379）、RabbitMQ（5672）和文件存储。")

    add_heading_custom(doc, "5.3 后端部署", level=2)
    add_paragraph_custom(doc, "（1）使用Maven构建项目：")
    add_paragraph_custom(doc, "cd career-planning-backend", indent=False)
    add_paragraph_custom(doc, "mvn clean package -DskipTests", indent=False)
    add_paragraph_custom(doc, "（2）配置application-prod.yaml：数据库连接、Redis连接、RabbitMQ连接、JWT密钥、AI服务URL")
    add_paragraph_custom(doc, "（3）创建systemd服务文件：/etc/systemd/system/career-backend.service")
    add_paragraph_custom(doc, "（4）启动服务：")
    add_paragraph_custom(doc, "sudo systemctl daemon-reload", indent=False)
    add_paragraph_custom(doc, "sudo systemctl enable career-backend", indent=False)
    add_paragraph_custom(doc, "sudo systemctl start career-backend", indent=False)

    add_heading_custom(doc, "5.4 前端部署", level=2)
    add_paragraph_custom(doc, "（1）安装依赖并构建：")
    add_paragraph_custom(doc, "cd career-planning-frontend", indent=False)
    add_paragraph_custom(doc, "npm install", indent=False)
    add_paragraph_custom(doc, "npm run build", indent=False)
    add_paragraph_custom(doc, "（2）配置Nginx：SSL证书、静态资源服务、API代理、Gzip压缩、SPA路由支持")
    add_paragraph_custom(doc, "（3）部署静态文件：")
    add_paragraph_custom(doc, "cp -r dist/* /var/www/career-front", indent=False)
    add_paragraph_custom(doc, "sudo nginx -t && sudo systemctl reload nginx", indent=False)

    add_heading_custom(doc, "5.5 AI服务部署", level=2)
    add_paragraph_custom(doc, "（1）Docker部署（推荐）：")
    add_paragraph_custom(doc, "docker build -t career-ai:latest .", indent=False)
    add_paragraph_custom(doc, "docker run -d --name career-ai -p 9000:9000 career-ai:latest", indent=False)
    add_paragraph_custom(doc, "（2）配置config.yaml：数据库连接、LLM模型配置（API Key）、日志路径、图数据库路径")

    add_heading_custom(doc, "5.6 监控与日志", level=2)
    add_paragraph_custom(doc, "（1）日志路径：")
    add_paragraph_custom(doc, "• /var/log/career-backend/：后端日志", indent=False)
    add_paragraph_custom(doc, "• /var/log/career-ai/：AI服务日志", indent=False)
    add_paragraph_custom(doc, "• /var/log/nginx/：Nginx日志", indent=False)
    add_paragraph_custom(doc, "（2）健康检查：")
    add_paragraph_custom(doc, "• 后端：curl http://localhost:8080/actuator/health", indent=False)
    add_paragraph_custom(doc, "• AI服务：curl http://localhost:9000/", indent=False)
    add_paragraph_custom(doc, "（3）监控建议：Prometheus + Grafana、Spring Boot Actuator、ELK Stack")

    add_heading_custom(doc, "5.7 备份与恢复", level=2)
    add_paragraph_custom(doc, "（1）数据库备份：每日凌晨2点自动备份，保留最近7天的备份")
    add_paragraph_custom(doc, "（2）恢复数据：解压备份文件后使用mysql命令恢复")

    # 保存文档
    output_path = "职引未来_技术文档_v4.0.docx"
    doc.save(output_path)
    print(f"文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_doc_from_content()
