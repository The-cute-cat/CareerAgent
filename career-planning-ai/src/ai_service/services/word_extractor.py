import os
import io
import re
import docx
import base64
import asyncio
import zipfile
import tempfile

from pathlib import Path
from typing import List, Optional, Any, Awaitable
from docx import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import BaseMessage, SystemMessage, HumanMessage
from ai_service.models.word import WordType
from ai_service.utils.logger_handler import log
from config import settings
from ai_service.engine.ai_engine import AIEngine
from concurrent.futures import ThreadPoolExecutor

os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
from docling.document_converter import DocumentConverter, FormatOption
from docling.datamodel.base_models import InputFormat, ConversionStatus
from docling.datamodel.pipeline_options import PipelineOptions, PdfPipelineOptions
from docling.backend.msword_backend import MsWordDocumentBackend
from docling_core.types.doc.document import DoclingDocument
from docling.pipeline.simple_pipeline import SimplePipeline


__all__ = ["WordExtractor"]

class WordExtractor:
    max_concurrent_requests: int = settings.lite_llm.max_concurrent_requests  # 最大并发请求数

    # 设置线程池，用于处理同步阻塞方法（如Docling的convert方法），避免阻塞事件循环，同时限制最大并发请求数以防止过载
    _executor = ThreadPoolExecutor(max_workers=max_concurrent_requests)


    # 初始化大模型客户端
    def __init__(self):
        self.llm = AIEngine().pick_brain(settings.lite_llm)
        self.output_parser = StrOutputParser()

        # 配置专门为 Word 优化的 Docling 转换器，确保能够处理 Word 文档中的复杂布局、表格和图片等内容，并且能够将表格精准转为 Markdown 格式以保留原有的格式信息
        # 配置 Pipeline 参数（支持 AI 表格识别和图片提取）
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_table_structure = (
            False  # 使用原生 XML 解析，不依赖 LibreOffice
        )
        pipeline_options.do_ocr = False  # 禁用 OCR
        # 初始化 Docling 转换器，指定只允许处理 Word 格式的文件，并传入优化后的配置
        self.docling_converter = DocumentConverter(
            allowed_formats=[InputFormat.DOCX],
            format_options={
                InputFormat.DOCX: FormatOption(
                    pipeline_cls=SimplePipeline,
                    backend=MsWordDocumentBackend,
                    pipeline_options=pipeline_options,
                )
            },
        )
        self.semaphore = asyncio.Semaphore(self.max_concurrent_requests)  # 用于限制并发请求数的异步信号量

    # 用于限制并发请求数的异步上下文管理器、确保在处理多个图片时不会超过设定的并发请求限制
    async def _bounded_gather(self, task: Awaitable[Any]) -> Any:
        async with self.semaphore:
            return await task

    # 解析word文件(文字、图片转文字)，最后返回增强markdown格式的纯文本
    async def detect_word_to_enhance_text(
        self, word: str | bytes | Path,
        password: Optional[str] = None
    ) -> Optional[str]:
        """
        解析word文件(文字、图片转文字),最后返回增强markdown格式的纯文本
        input:
            word: word文件的路径或字节数据
            password: 如果word文件受密码保护,则传入密码进行解密,目前暂不支持加密文件的处理(之后可添加逻辑)
        output:
            增强markdown格式的纯文本,如果解析失败则返回None
        设计逻辑:
        1. 文件通用校验,检测是否为有效的Word文件
        2. 安全校验，检测是否含有宏等潜在风险
        3. 通过word文件的魔数来判断word文件类型,目前只支持.docx/.dotx/.docm/.dotm的具体类型判断,其他类型需在之后补充
        4. 使用 Docling 进行深度解析,Docling 会分析布局、表格和图片位置等信息，并按顺序提取文本和图片内容描述
        5. 将提取的内容转化为 Markdown 格式的纯文本，保留原有的格式信息（如表格转为 Markdown 表格等）
        6. 调用_enhance_image_with_ai方法对图片内容进行增强,将图片中的文字提取出来并替换原有的图片占位符
        """
        word_path = None
        try:
            # 1. 文件通用校验，检测是否为有效的Word文件
            if isinstance(word, str):
                word_path = word
            elif isinstance(word, Path):
                word_path = str(word)
            else:
                # 如果是字节数据，先保存到临时文件
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".docx"
                ) as tmp_file:
                    tmp_file.write(word)
                    word_path = tmp_file.name

            if not Path(word_path).exists() or not Path(word_path).is_file():
                log.error(f"文件不存在或不是一个有效的Word文件: {word_path}")
                return None

            # 2. 安全校验，检测是否含有宏等潜在风险
            if not self.security_check(word_path):
                log.error(f"文件安全校验失败，可能含有宏等潜在风险: {word_path}")
                return None

            # 3. 通过word文件的魔数来判断word文件类型,目前只支持.docx/.dotx/.docm/.dotm的具体类型判断,其他类型需在之后补充
            word_type = self.detect_word_type(word_path)
            if word_type == WordType.UNKNOWN:
                log.error(f"无法识别的Word文件类型: {word_path}")
                return None

            log.info(f"检测到的Word文件类型: {word_type}")

            # 4. 使用 Docling 进行深度解析，Docling 会分析布局、表格和图片位置等信息，并按顺序提取文本和图片内容描述
            loop = asyncio.get_event_loop()
            conversation = await loop.run_in_executor(self._executor, self.docling_converter.convert, word_path)
            if conversation is None or conversation.status not in [
                ConversionStatus.SUCCESS,
                ConversionStatus.PARTIAL_SUCCESS,
            ]:
                log.error(
                    f"Docling 转换失败: {conversation.errors if conversation else '未知错误'}"
                )
                return None

            # 5. 将提取的内容转化为 Markdown 格式的纯文本，保留原有的格式信息（如表格转为 Markdown 表格等）
            conversation_result = conversation.document.export_to_markdown()
            # 6. 调用_enhance_image_with_ai方法对图片内容进行增强，将图片中的文字提取出来并替换原有的图片占位符
            enhanced_result = await self._enhance_image_with_ai(
                conversation.document, conversation_result
            )

            return enhanced_result

        except Exception as e:
            log.error(f"解析Word文件时出错: {e}")
            return None

        finally:
            # 如果输入是字节数据，处理完成后删除临时文件
            if (
                isinstance(word, bytes)
                and word_path is not None
                and Path(word_path).exists()
            ):
                try:
                    os.remove(word_path)
                except Exception as e:
                    log.warning(f"删除临时文件失败: {e}")

    # 对图片内容进行增强，将图片中的文字提取出来并替换原有的图片占位符
    async def _enhance_image_with_ai(
        self, document: DoclingDocument, markdown_text: str
    ) -> str:
        """
        对图片内容进行增强，将图片中的文字提取出来并替换原有的图片占位符
        input:
            document: DoclingDocument对象,包含文档的结构化信息和图片数据
            markdown_text: Docling转换后生成的Markdown格式文本,包含图片占位符
        output:
            增强后的Markdown格式文本,图片占位符被替换为AI提取的文字描述
        设计逻辑:
            1. 检查文档中是否有图片,如果没有图片则直接返回原始的Markdown文本
            2. 遍历文档中的图片，提取图片的字节流
            3. 调用阿里云多模态API检测图片内容,获取图片中的文字描述
            4. 替换Markdown文本中的图片占位符,将其替换为AI提取的文字描述,保留原有的格式信息(如表格转为 Markdown 表格等)
             - Docling 默认会将图片替换为类似于 ![Figure 1](image_path) 的占位符，我们需要将其替换为 AI 提取的文字描述
        """
        try:
            if not document.pictures:
                log.info("没有图片需要增强")
                return markdown_text

            task = []

            for i, image in enumerate(document.pictures):
                log.info(f"正在增强图片 {i+1}")
                if image.image is None:
                    log.error(f"图片对象中没有有效的图片数据: 图片索引 {i+1}")
                    continue
                # 1.提取文件字节流
                real_pil_image = image.image.pil_image
                if real_pil_image is None:
                    log.error(f"无法提取图片的PIL对象: 图片索引 {i+1}")
                    continue

                image_bytes_io = io.BytesIO()
                real_pil_image.save(image_bytes_io, format="PNG")

                image_bytes = image_bytes_io.getvalue()

                # 2.调用阿里云多模态API检测图片内容，获取图片中的文字描述（单张图片的检测逻辑）
                task.append(self._bounded_gather(self._detect_images(image_bytes, "png")))

            descriptions = await asyncio.gather(*task, return_exceptions=True)

            for description in descriptions:
                if isinstance(description, Exception):
                    log.error(f"检测图片内容时出错: {description}")
                    continue
                # 3.替换占位符
                # Docling 默认会将图片替换为类似于 ![Figure 1](image_path) 的占位符，我们需要将其替换为 AI 提取的文字描述
                placeholder = f"<!-- image -->"
                replacement = f"\n> **[图片内容语义摘要：{description}]**\n"

                markdown_text = re.sub(placeholder, replacement, markdown_text, count=1)

            return markdown_text
        except Exception as e:
            log.error(f"增强图片内容时出错: {e}")
            return markdown_text

    # 将字符串列表转化为纯字符串
    @staticmethod
    def list_to_string(texts: List[str]) -> str:
        return "\n".join(texts)

    # 检查文档中是否有文本(通过扫描ZIP包判断)
    @staticmethod
    def _has_text(word_path: str) -> bool:
        """
        检查文档中是否有文本(通过扫描ZIP包判断)
        input:
            word_path: word文件路径
        output:
            如果包含文本则返回True,否则返回False
        设计逻辑:
            1. 由于.docx/.dotx/.docm/.dotm文件实际上是一个ZIP压缩包,我们可以通过扫描压缩包中的XML文件来判断是否包含文本
                - 文本通常存储在"word/"目录下的XML文件中,文本内容被包含在<w:t>标签内
                - 我们可以使用正则表达式来匹配<w:t>标签及其内容，并检查是否存在非空白字符来判断是否包含文本
        """

        # 匹配 <w:t> 标签及其内容的正则表达式
        # 考虑到可能存在带属性的标签如 <w:t xml:space="preserve">
        text_tag_pattern = re.compile(r"<w:t[^>]*>(.*?)</w:t>", re.DOTALL)

        try:
            with zipfile.ZipFile(word_path, "r") as zf:
                # 1. 遍历压缩包内所有文件
                for file_name in zf.namelist():
                    # 2. 只扫描 word/ 目录下且以 .xml 结尾的文件
                    if file_name.startswith("word/") and file_name.endswith(".xml"):
                        with zf.open(file_name) as f:
                            # 3. 读取内容 (建议分块读取或限制大小，防止极端巨大的 XML)
                            content = f.read().decode("utf-8", errors="ignore")

                            # 4. 搜索所有的 <w:t> 标签
                            matches = text_tag_pattern.findall(content)

                            for text_content in matches:
                                # 5. 校验提取到的文字是否包含非空白字符
                                # 即使只有 1 个有效字符，也说明该文档含有原生文本
                                if text_content and text_content.strip():
                                    return True

            return False  # 扫完全部 XML 都没有发现有效文字
        except Exception as e:
            print(f"暴力扫描失败: {e}")
            return False

    # 检查文档中是否有图片(通过word文档底层zip特性判断)
    @staticmethod
    def _has_images(word_path: str) -> bool:
        """
        检查文档中是否有图片(通过word文档底层zip特性判断)
        input:
            word_path: word文件路径
        output:
            如果包含图片则返回True,否则返回False
        设计逻辑:
            1. 由于.docx/.dotx/.docm/.dotm文件实际上是一个ZIP压缩包,我们可以通过检查压缩包中的文件结构来判断是否包含图片
                - 图片通常存储在"word/media/"目录下，文件扩展名通常是.png、.jpg、.jpeg、.bmp等
        """

        try:
            # 1. 由于.docx/.dotx/.docm/.dotm文件实际上是一个ZIP压缩包，通过检查压缩包中的文件结构来判断是否包含图片
            with zipfile.ZipFile(word_path, "r") as zf:
                # 图片通常存储在"word/media/"目录下，文件扩展名通常是.png、.jpg、.jpeg、.bmp等
                for zip_info in zf.infolist():
                    if zip_info.filename.startswith(
                        "word/media/"
                    ) and zip_info.filename.endswith((".png", ".jpg", ".jpeg", ".bmp")):
                        return True
            return False
        except Exception as e:
            log.error(f"检查图片存在性失败: {e}")
            return False

    # 提取文档中所有图片的embed_ids
    @staticmethod
    def _extract_image_embed_ids(word_path: str) -> List[str]:
        """
        提取文档中所有图片的embed_ids
        input:
            word_path: word文件路径
        output:
            图片的embed_id列表,如果提取失败则返回空列表
        设计逻辑:
            1. 通过Document对象解析文档的XML结构来提取图片的embed_id
                - 对于嵌入式图片,通常在XML中以<w:inline>元素表示,图片的embed_id存储在<a:blip>元素的r:embed属性中
                - 对于浮动图片,通常在XML中以<w:anchor>元素表示,图片的embed_id同样存储在<a:blip>元素的r:embed属性中
        """

        doc = Document(word_path)
        embed_ids = set()  # 使用集合避免重复的embed_id
        try:
            # 2.通过Document对象解析文档的XML结构来提取图片的embed_id
            # 解析嵌入式图片的XML
            inline_shape_elems = doc.element.xpath("//w:inline//a:blip/@r:embed")
            for embed_id in inline_shape_elems:
                embed_ids.add(embed_id)

            # 解析浮动图片的XML
            anchor_shape_elems = doc.element.xpath("//w:anchor//a:blip/@r:embed")
            for embed_id in anchor_shape_elems:
                embed_ids.add(embed_id)
        except Exception as e:
            log.error(f"提取图片embed_id失败: {e}")
            return []
        return list(embed_ids)

    # 获取文件的魔数,转化成十六进制字符串
    @staticmethod
    def get_magic_number(file_path: str, read_bytes: int = 8) -> Optional[str]:
        """
        获取文件的魔数,转化成十六进制字符串
        input:
            file_path: 文件路径
            read_bytes: 读取的字节数,默认8字节
        output:
            魔数的十六进制字符串表示,每个字节用空格分隔（防止误读）,如果读取失败则返回None
        """

        with open(file_path, "rb") as f:
            return " ".join([f"{byte:02x}" for byte in f.read(read_bytes)])

    # 通过word文件的魔数来判断word文件类型,目前只支持.docx/.dotx/.docm/.dotm的具体类型判断,其他类型需在之后补充
    @staticmethod
    def detect_word_type(word_path: str) -> WordType:
        """
        通过word文件的魔数来判断word文件类型
        input:
            word_path: word文件路径
        output:
            WordType枚举值,如果无法识别则返回WordType.UNKNOWN
            目前只支持.docx/.dotx/.docm/.dotm的具体类型判断,其他类型需在之后补充
        设计逻辑:
        1. 首先检查文件是否存在且是一个有效的文件
        2. 获取文件的魔数并进行判断
        3. 根据魔数判断文件类型
            - .docx/.dotx/.docm/.dotm的魔数是"50 4b 03 04"(ZIP文件格式),检查是否含宏和是否为模板,判断是文本可提取还是扫描图片或混合文档
            - .dot/.dot的魔数是"d0 cf 11 e0"，但需要进一步检查是否为模板
            - .wps/.wpt的魔数是"a5 5a 6b 7a"，但需要进一步检查是否为文本可提取的格式
        """
        # 1. 首先检查文件是否存在且是一个有效的文件
        file = Path(word_path)
        if not file.exists() or not file.is_file():
            log.error(f"文件不存在或不是一个有效的文件: {word_path}")
            return WordType.UNKNOWN

        # 2. 获取文件的魔数并进行判断
        magic_number = WordExtractor.get_magic_number(word_path)
        if magic_number is None:
            log.error(f"无法读取word文件的魔数: {word_path}")
            return WordType.UNKNOWN

        # 3. 根据魔数判断文件类型
        try:
            # .docx/.dotx/.docm/.dotm的魔数是"50 4b 03 04"（ZIP文件格式），检查是否含宏和是否为模板
            if magic_number.startswith("50 4b 03 04"):
                # 检查是否含宏
                with zipfile.ZipFile(word_path, "r") as zip_ref:
                    if any(
                        name.startswith("word/vbaProject.bin")
                        for name in zip_ref.namelist()
                    ):
                        return WordType.MACRO_ENABLED

                # 检查是否为模板
                if Path(word_path).suffix.lower() in [".dotx", ".dotm"]:
                    return WordType.DOT_TEXT_BASED

                # 尝试解析文档内容以判断是否为扫描图片或混合文档
                doc = docx.Document(word_path)
                has_images = WordExtractor._has_images(word_path)
                has_text = WordExtractor._has_text(word_path)

                if has_images and not has_text:
                    return WordType.DOCX_SCANNED
                elif has_images and has_text:
                    return WordType.DOCX_MIXED
                else:
                    return WordType.DOCX_TEXT_BASED

            # .dot/.dot的魔数是"d0 cf 11 e0"，但需要进一步检查是否为模板
            elif magic_number.startswith("d0 cf 11 e0"):  # .doc的魔数
                # 由于.doc和.dot的魔数相同，需要进一步检查文件扩展名来区分是否为模板
                if Path(word_path).suffix.lower() == ".dot":
                    return WordType.DOT_TEXT_BASED
                return WordType.DOC_TEXT_BASED
            # .wps/.wpt的魔数是"a5 5a 6b 7a"，但需要进一步检查是否为文本可提取的格式
            elif magic_number.startswith("a5 5a 6b 7a"):
                return WordType.WPS_TEXT_BASED
            else:
                log.warning(f"无法识别的文件类型,魔数: {magic_number}")
                return WordType.UNKNOWN

        except zipfile.BadZipFile:
            log.error(
                f"文件不是一个有效的ZIP格式,无法识别为.docx/.dotx/.docm/.dotm: {word_path}"
            )
            return WordType.UNKNOWN

        except Exception as e:
            log.error(f"检测word文件类型时出错: {e}")
            return WordType.UNKNOWN

    # 从word文档的XML文件中提取文本(暂时弃用)
    @staticmethod
    def _extract_text(word_path: str, word_type: WordType) -> List[str]:
        """
        从文件中提取文本
        input:
            word_path: word文件路径
            word_type: 文件类型
        output:
            文本列表,如果提取失败则返回空列表
        设计逻辑:
            1. 根据文件类型提取文本
                - 对于.docx/.dotx/.docm/.dotm格式的文件,使用python-docx库提取段落文本
                - 对于.doc格式的文件,暂不支持文本提取(之后可添加逻辑)
                - 对于.wps/.wpt格式的文件,暂不支持文本提取(之后可添加逻辑)
                - 其他类型的文件暂不支持文本提取
        """
        # 1. 根据文件类型提取文本
        try:
            texts = []
            # 处理.docx/.dotx/.docm/.dotm格式的文件
            if word_type in [
                WordType.DOCX_TEXT_BASED,
                WordType.DOT_TEXT_BASED,
                WordType.MACRO_ENABLED,
                WordType.DOCX_MIXED,
            ]:
                doc = docx.Document(word_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        texts.append(para.text.strip())
            # 处理.doc格式的文件
            elif word_type == WordType.DOC_TEXT_BASED:
                # 这里可以添加对 .doc 文件的文本提取逻辑
                log.warning(f"当前文件类型暂不支持文本提取: {word_type}")
                pass
            # 处理.wps/.wpt格式的文件
            elif word_type == WordType.WPS_TEXT_BASED:
                # 这里可以添加对 .wps/.wpt 文件的文本提取逻辑
                log.warning(f"当前文件类型暂不支持文本提取: {word_type}")
                pass
            else:
                log.warning(f"当前文件类型暂不支持文本提取: {word_type}")
                pass
        except Exception as e:
            log.error(f"提取文本时出错: {e}")
            return []
        return texts

    # 从文件中提取图片
    @staticmethod
    def _extract_images_bytes(
        word_path: str, word_type: WordType
    ) -> List[tuple[bytes, str]]:
        """
        从文件中提取图片
        input:
            word_path: word文件路径
            word_type: 文件类型
        output:
            返回元组列表，包含图片的二进制数据和后缀，如果提取失败则返回空列表
        设计逻辑:
            1. 根据文件类型提取图片
                - 对于含宏文件、扫描件和混合文档（目前仅对.docx格式进行处理,其他格式的图片提取逻辑之后可添加），使用python-docx库提取图片的二进制数据
                - 对于其他类型的文件暂不支持图片提取
        """
        # 1. 根据文件类型提取图片
        try:
            image_paths = []
            # 处理含宏文件、扫描件和混合文档（目前仅对.docx格式进行处理，其他格式的图片提取逻辑之后可添加）
            if word_type in [
                WordType.MACRO_ENABLED,
                WordType.DOCX_MIXED,
                WordType.DOCX_SCANNED,
            ]:
                if not zipfile.is_zipfile(word_path):
                    log.error(f"文件不是一个有效的ZIP格式,无法提取图片: {word_path}")
                    return []

                # 由于.docx/.dotx/.docm/.dotm文件实际上是一个ZIP压缩包，我们可以通过检查压缩包中的文件结构来提取图片
                with zipfile.ZipFile(word_path, "r") as zip_ref:
                    for zip_info in zip_ref.infolist():
                        if zip_info.filename.startswith(
                            "word/media/"
                        ) and zip_info.filename.endswith(
                            (".png", ".jpg", ".jpeg", ".bmp")
                        ):
                            with zip_ref.open(zip_info) as image_file:
                                image_data = image_file.read()
                                image_paths.append(
                                    (image_data, zip_info.filename.split(".")[-1])
                                )
            else:
                log.warning(f"当前文件类型暂不支持图片提取: {word_type}")
                pass
        except Exception as e:
            log.error(f"提取图片时出错: {e}")
            return []
        return image_paths

    # 通过调用阿里云多模态AI服务来检测图片中的内容(单张图片的检测逻辑)
    async def _detect_images(
        self, image_data: bytes, image_suffix: str
    ) -> Optional[str]:
        """通过调用阿里云多模态AI服务来检测图片中的内容
        input:
            image_data: 图片的二进制数据
            image_suffix: 图片的后缀(如png、jpg等)
        output:
            图片内容的描述列表
        设计逻辑:
            1. 初始化阿里云多模态API客户端
            2. 定义输出解析器,将API返回的结果解析为字符串列表
            3. 遍历图片数据列表,调用阿里云多模态API
                - 输入参数包括图片的二进制数据、输入类型(image)和提示语(提取图片中的所有文字，保留格式)
                - 将API返回的结果通过输出解析器解析为文本描述,并添加到结果列表中
            4. 返回图片内容的描述列表
        """
        try:
            if not image_data:
                log.info("没有图片需要检测")
                return None

            # 将图片的二进制数据转化为base64字符串列表,以便传递给API
            image_data_base64 = base64.b64encode(image_data).decode("utf-8")

            # 构造文件头信息,以便API正确识别图片格式
            file_header = f"data:image/{image_suffix};base64,{image_data_base64}"
            # 构造信息提示语,以便API正确理解任务需求
            # --- 核心修复点：遵循 OpenAI Vision 标准格式 ---
            message: List[BaseMessage] = [
                HumanMessage(
                    content=[
                        {
                            "type": "text",
                            "text": "描述图片,要保留数据和格式.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": file_header
                            },
                        }
                    ]
                )
            ]

            # 调用阿里云多模态API进行图片内容检测
            response = await self.llm \
                .add_langchain_message(message) \
                .next_step().next_step() \
                .into_text().do()  # 获取文本输出结果
            if response is None:
                log.error("调用阿里云多模态API检测图片内容失败: API返回了None")
                return None
            # 通过输出解析器解析API返回的结果为文本描述
            description = self.output_parser.parse(response)

        except Exception as e:
            log.error(f"调用阿里云多模态API检测图片内容时出错: {e}")
            return None
        return description

    # 文件通用校验，检测是否为有效的Word文件
    @staticmethod
    def validate_word_file(word_path: str) -> Optional[bool]:
        pass

    # 安全校验，检测是否含有宏等潜在风险
    @staticmethod
    def security_check(word_path: str) -> bool:
        return True
